"""
微信聊天记录渲染器

将聊天消息渲染为仿微信对话界面的 HTML 页面或 Word 文档段落，
用于生成可直接附入证据材料的聊天记录展示。
"""

from __future__ import annotations

import base64
import html as html_lib
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_WECHAT_CSS = """\
* { margin: 0; padding: 0; box-sizing: border-box; }
body {
    background: #EBEBEB; font-family: -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif;
    font-size: 15px; color: #000;
}
.chat-container { max-width: 420px; margin: 0 auto; background: #EBEBEB; min-height: 100vh; }
.chat-header {
    background: #EDEDED; text-align: center; padding: 12px 0 10px;
    border-bottom: 1px solid #D6D6D6; font-size: 17px; font-weight: 500;
}
.msg-row { display: flex; padding: 10px 12px; align-items: flex-start; }
.msg-row.sent { flex-direction: row-reverse; }
.avatar {
    width: 40px; height: 40px; border-radius: 4px; flex-shrink: 0;
    display: flex; align-items: center; justify-content: center;
    font-size: 16px; font-weight: 600; color: #fff;
}
.avatar.received { background: #7BB1EA; }
.avatar.sent     { background: #6BCB77; }
.bubble-wrap { max-width: 260px; margin: 0 8px; }
.bubble {
    padding: 9px 12px; border-radius: 6px; line-height: 1.5;
    word-break: break-all; position: relative; font-size: 15px;
}
.bubble.received { background: #fff; }
.bubble.sent     { background: #95EC69; }
.timestamp {
    text-align: center; color: #999; font-size: 12px; padding: 8px 0 2px;
}
.transfer-box {
    background: #FA9D3B; color: #fff; padding: 10px 14px;
    border-radius: 6px; min-width: 180px;
}
.transfer-box .amount { font-size: 18px; font-weight: 600; }
.transfer-box .label  { font-size: 12px; margin-top: 4px; }
.voice-box {
    display: flex; align-items: center; gap: 6px;
    padding: 9px 14px; border-radius: 6px; min-width: 80px;
}
.voice-box.received { background: #fff; }
.voice-box.sent     { background: #95EC69; }
.voice-icon { font-size: 18px; }
.voice-dur  { font-size: 13px; color: #666; }
.img-msg img { max-width: 200px; border-radius: 4px; }
"""

_HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>微信聊天记录 - {contact}</title>
<style>
{css}
</style>
</head>
<body>
<div class="chat-container">
<div class="chat-header">{contact}</div>
{messages_html}
</div>
</body>
</html>
"""


class ChatRenderer:
    """将聊天消息列表渲染为仿微信界面的 HTML 或 Word 格式。

    消息字典结构::

        {
            "sender": "张三",          # 发送者名称
            "type": "text",            # text / image / voice / transfer
            "content": "你好",         # 文本内容或图片路径
            "timestamp": datetime(...),
            "duration": 5,             # 语音时长(秒)，仅 voice 类型
            "amount": 500.00,          # 转账金额，仅 transfer 类型
        }
    """

    # ------------------------------------------------------------------ #
    #  HTML 渲染
    # ------------------------------------------------------------------ #

    @staticmethod
    def render_chat_html(
        messages: List[Dict[str, Any]],
        contact_name: str,
        my_name: str = "我",
    ) -> str:
        """将消息列表渲染为仿微信对话 HTML 页面。

        Args:
            messages: 消息列表。
            contact_name: 联系人名称。
            my_name: 当前用户名称，默认"我"。

        Returns:
            完整的 HTML 字符串。
        """
        parts: List[str] = []
        last_ts: Optional[datetime] = None

        for msg in messages:
            ts: Optional[datetime] = msg.get("timestamp")
            # 每隔 5 分钟显示一次时间戳
            if ts and (last_ts is None or (ts - last_ts).total_seconds() > 300):
                parts.append(
                    f'<div class="timestamp">{ts.strftime("%Y-%m-%d %H:%M")}</div>'
                )
                last_ts = ts

            sender = msg.get("sender", "")
            is_sent = sender == my_name
            direction = "sent" if is_sent else "received"
            initials = sender[:1] if sender else "?"

            bubble_html = ChatRenderer._render_bubble(msg, direction)

            parts.append(
                f'<div class="msg-row {direction}">'
                f'<div class="avatar {direction}">{html_lib.escape(initials)}</div>'
                f'<div class="bubble-wrap">{bubble_html}</div>'
                f"</div>"
            )

        return _HTML_TEMPLATE.format(
            contact=html_lib.escape(contact_name),
            css=_WECHAT_CSS,
            messages_html="\n".join(parts),
        )

    @staticmethod
    def _render_bubble(msg: Dict[str, Any], direction: str) -> str:
        """根据消息类型渲染气泡 HTML。"""
        msg_type = msg.get("type", "text")
        content = msg.get("content", "")

        if msg_type == "image":
            img_data = ChatRenderer._embed_image(content)
            return f'<div class="img-msg"><img src="{img_data}" alt="图片"></div>'

        if msg_type == "voice":
            dur = msg.get("duration", 0)
            return (
                f'<div class="voice-box {direction}">'
                f'<span class="voice-icon">🎤</span>'
                f'<span class="voice-dur">{dur}″</span>'
                f"</div>"
            )

        if msg_type == "transfer":
            amount = msg.get("amount", 0)
            return (
                f'<div class="transfer-box">'
                f'<div class="amount">¥{amount:.2f}</div>'
                f'<div class="label">微信转账</div>'
                f"</div>"
            )

        # 默认 text
        return f'<div class="bubble {direction}">{html_lib.escape(str(content))}</div>'

    @staticmethod
    def _embed_image(path_or_url: str) -> str:
        """将本地图片转为 base64 data URI；若不存在则返回原路径。"""
        p = Path(path_or_url)
        if p.is_file():
            mime = mimetypes.guess_type(str(p))[0] or "image/png"
            data = base64.b64encode(p.read_bytes()).decode()
            return f"data:{mime};base64,{data}"
        return html_lib.escape(path_or_url)

    # ------------------------------------------------------------------ #
    #  HTML → 图片（需要 playwright / selenium）
    # ------------------------------------------------------------------ #

    @staticmethod
    def render_to_image(html_content: str, output_path: str) -> Path:
        """将 HTML 渲染为图片。

        注意：完整的截图功能需要安装 playwright 或 selenium 并配置
        headless 浏览器。当前实现仅将 HTML 保存到磁盘。

        Args:
            html_content: 完整 HTML 字符串。
            output_path: 输出路径（建议 .html 后缀）。

        Returns:
            保存后的文件 Path。
        """
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)

        # 尝试使用 playwright（若已安装）
        try:
            from playwright.sync_api import sync_playwright  # type: ignore

            html_tmp = out.with_suffix(".tmp.html")
            html_tmp.write_text(html_content, encoding="utf-8")

            with sync_playwright() as pw:
                browser = pw.chromium.launch(headless=True)
                page = browser.new_page(viewport={"width": 460, "height": 800})
                page.goto(f"file://{html_tmp.resolve()}")
                page.wait_for_load_state("networkidle")
                page.screenshot(path=str(out.with_suffix(".png")), full_page=True)
                browser.close()

            html_tmp.unlink(missing_ok=True)
            return out.with_suffix(".png")
        except ImportError:
            # playwright 未安装，回退为保存 HTML
            out = out.with_suffix(".html")
            out.write_text(html_content, encoding="utf-8")
            return out

    # ------------------------------------------------------------------ #
    #  Word 文档内嵌渲染
    # ------------------------------------------------------------------ #

    @staticmethod
    def render_chat_to_docx(
        messages: List[Dict[str, Any]],
        contact_name: str,
        doc: Document,
        my_name: str = "我",
    ) -> None:
        """将聊天消息渲染为 Word 文档段落。

        在传入的 Document 末尾追加格式化的聊天记录，用缩进和颜色
        区分收发双方，每条消息包含时间、发送者和内容。

        Args:
            messages: 消息列表。
            contact_name: 联系人名称。
            doc: python-docx Document 对象（会被就地修改）。
            my_name: 当前用户名称。
        """
        # 小标题
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(f"与 {contact_name} 的聊天记录")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "黑体"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        doc.add_paragraph()  # 空行

        last_ts: Optional[datetime] = None

        for msg in messages:
            ts: Optional[datetime] = msg.get("timestamp")
            sender = msg.get("sender", "")
            is_sent = sender == my_name
            msg_type = msg.get("type", "text")
            content = msg.get("content", "")

            # 时间戳分隔
            if ts and (last_ts is None or (ts - last_ts).total_seconds() > 300):
                ts_p = doc.add_paragraph()
                ts_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ts_run = ts_p.add_run(ts.strftime("%Y-%m-%d %H:%M"))
                ts_run.font.size = Pt(9)
                ts_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                ts_run.font.name = "宋体"
                ts_run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                last_ts = ts

            # 消息段落
            p = doc.add_paragraph()
            if is_sent:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.left_indent = Cm(4)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.right_indent = Cm(4)

            # 发送者
            name_run = p.add_run(f"{sender}：")
            name_run.bold = True
            name_run.font.size = Pt(10)
            name_run.font.name = "宋体"
            name_run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if is_sent:
                name_run.font.color.rgb = RGBColor(0x07, 0xC1, 0x60)
            else:
                name_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            # 内容
            content_text = ChatRenderer._format_content_for_docx(msg_type, msg)
            c_run = p.add_run(content_text)
            c_run.font.size = Pt(10.5)
            c_run.font.name = "宋体"
            c_run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    @staticmethod
    def _format_content_for_docx(msg_type: str, msg: Dict[str, Any]) -> str:
        """将不同类型的消息格式化为纯文本。"""
        content = msg.get("content", "")
        if msg_type == "image":
            return "[图片]"
        if msg_type == "voice":
            dur = msg.get("duration", 0)
            return f"[语音 {dur}秒]"
        if msg_type == "transfer":
            amount = msg.get("amount", 0)
            return f"[微信转账] ¥{amount:.2f}"
        return str(content)
