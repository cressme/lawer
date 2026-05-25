"""
证据清单生成器

生成符合中国法院要求的证据清单Word文档，包含：
- 标准证据清单表头（案由、原告、被告等）
- 证据明细表格（序号、证据名称、来源、摘要、证明目的、页码、备注）
- 规范的中文排版（宋体正文、黑体标题）
"""

from __future__ import annotations

import os
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# 证据类别常量
EVIDENCE_CATEGORIES: List[str] = [
    "聊天记录",
    "转账记录",
    "图片",
    "语音",
    "视频",
    "文件",
]

# 默认分类排序权重（序号越小越靠前）
_CATEGORY_ORDER: Dict[str, int] = {cat: i for i, cat in enumerate(EVIDENCE_CATEGORIES)}


@dataclass
class EvidenceItem:
    """单条证据项。

    Attributes:
        id: 证据编号（可由 auto_number_evidence 自动赋值）。
        category: 证据类别，如"聊天记录"、"转账记录"等。
        title: 证据名称 / 标题。
        description: 证据内容摘要。
        source: 证据来源说明（如"微信PC客户端本地数据"）。
        timestamp: 证据产生时间。
        content: 文本内容或媒体文件路径。
        proof_purpose: 证明目的。
        page_range: 页码范围，如"1-3"。
        remarks: 备注。
    """

    id: Optional[int] = None
    category: str = "聊天记录"
    title: str = ""
    description: str = ""
    source: str = ""
    timestamp: Optional[datetime] = None
    content: str = ""
    proof_purpose: str = ""
    page_range: str = ""
    remarks: str = ""


def auto_number_evidence(evidence_items: List[EvidenceItem]) -> List[EvidenceItem]:
    """按类别分组后自动分配连续序号。

    排序规则：先按 EVIDENCE_CATEGORIES 中的预定义顺序，再按时间戳升序。
    序号从 1 开始连续编排。

    Args:
        evidence_items: 待编号的证据列表。

    Returns:
        编号后的证据列表（原列表不变，返回新列表）。
    """
    # 按类别分组
    groups: Dict[str, List[EvidenceItem]] = defaultdict(list)
    for item in evidence_items:
        groups[item.category].append(item)

    # 每组内按时间排序
    for cat_items in groups.values():
        cat_items.sort(key=lambda x: x.timestamp or datetime.min)

    # 按类别权重排序后合并
    sorted_categories = sorted(
        groups.keys(),
        key=lambda c: _CATEGORY_ORDER.get(c, len(EVIDENCE_CATEGORIES)),
    )

    numbered: List[EvidenceItem] = []
    seq = 1
    for cat in sorted_categories:
        for item in groups[cat]:
            from dataclasses import replace

            numbered.append(replace(item, id=seq))
            seq += 1

    return numbered


class EvidenceListGenerator:
    """生成法院证据清单 Word 文档。

    用法示例::

        gen = EvidenceListGenerator()
        path = gen.generate(
            case_info={"case_type": "民间借贷纠纷", "plaintiff": "张三", ...},
            evidence_items=[...],
            output_path="./output/证据清单.docx",
        )
    """

    # 默认模板配置
    _DEFAULT_CONFIG: dict = {
        "body_font": "宋体",
        "header_font": "黑体",
        "body_size_pt": 12,
        "header_size_pt": 22,
        "sub_header_size_pt": 14,
        "table_font_size_pt": 10.5,
        "page_margin_cm": 2.54,
    }

    def __init__(self, template_config: Optional[dict] = None) -> None:
        """初始化生成器。

        Args:
            template_config: 可选的模板配置，覆盖默认值。
        """
        self.config: dict = {**self._DEFAULT_CONFIG, **(template_config or {})}

    # ------------------------------------------------------------------ #
    #  公开 API
    # ------------------------------------------------------------------ #

    def generate(
        self,
        case_info: dict,
        evidence_items: List[EvidenceItem],
        output_path: str,
    ) -> Path:
        """生成证据清单 Word 文档。

        Args:
            case_info: 案件信息字典，包含 case_type / plaintiff / defendant /
                       court_name / case_number 等键。
            evidence_items: 证据列表（如尚未编号则自动编号）。
            output_path: 输出文件路径。

        Returns:
            生成的文件 Path 对象。
        """
        # 自动编号
        items = auto_number_evidence(evidence_items)

        doc = Document()
        self._setup_page(doc)
        self._add_header(doc, case_info)
        self._add_evidence_table(doc, items)
        self._add_footer(doc, case_info)
        self._add_page_number(doc)

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return out

    # ------------------------------------------------------------------ #
    #  页面设置
    # ------------------------------------------------------------------ #

    def _setup_page(self, doc: Document) -> None:
        """设置页面边距与默认字体。"""
        section = doc.sections[0]
        margin = Cm(self.config["page_margin_cm"])
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

        # 默认正文样式
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.config["body_font"]
        font.size = Pt(self.config["body_size_pt"])
        style.element.rPr.rFonts.set(qn("w:eastAsia"), self.config["body_font"])

    # ------------------------------------------------------------------ #
    #  文档标题 + 案件信息
    # ------------------------------------------------------------------ #

    def _add_header(self, doc: Document, case_info: dict) -> None:
        """添加证据清单标题与案件基本信息。"""
        # 标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("证  据  清  单")
        run.bold = True
        run.font.size = Pt(self.config["header_size_pt"])
        run.font.name = self.config["header_font"]
        run.element.rPr.rFonts.set(qn("w:eastAsia"), self.config["header_font"])

        # 案件信息行
        court_name = case_info.get("court_name", "")
        case_number = case_info.get("case_number", "")
        if court_name or case_number:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub.add_run(f"{court_name}  {case_number}")
            sub_run.font.size = Pt(self.config["sub_header_size_pt"])
            sub_run.font.name = self.config["body_font"]
            sub_run.element.rPr.rFonts.set(qn("w:eastAsia"), self.config["body_font"])

        info_lines = [
            ("案　　由", case_info.get("case_type", "")),
            ("原　　告", case_info.get("plaintiff", "")),
            ("被　　告", case_info.get("defendant", "")),
        ]
        for label, value in info_lines:
            p = doc.add_paragraph()
            run_label = p.add_run(f"{label}：")
            run_label.bold = True
            self._set_run_font(run_label, self.config["body_font"], self.config["body_size_pt"])
            run_value = p.add_run(value)
            self._set_run_font(run_value, self.config["body_font"], self.config["body_size_pt"])

        # 空行
        doc.add_paragraph()

    # ------------------------------------------------------------------ #
    #  证据明细表格
    # ------------------------------------------------------------------ #

    _TABLE_HEADERS: List[str] = [
        "序号", "证据名称", "证据来源", "证据内容摘要", "证明目的", "页码", "备注",
    ]

    _COL_WIDTHS_CM: List[float] = [1.2, 3.0, 2.5, 4.5, 3.5, 1.5, 1.8]

    def _add_evidence_table(self, doc: Document, items: List[EvidenceItem]) -> None:
        """添加证据明细表格。"""
        num_cols = len(self._TABLE_HEADERS)
        table = doc.add_table(rows=1 + len(items), cols=num_cols, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # 列宽
        for idx, width_cm in enumerate(self._COL_WIDTHS_CM):
            for row in table.rows:
                row.cells[idx].width = Cm(width_cm)

        # 表头行
        header_row = table.rows[0]
        for idx, text in enumerate(self._TABLE_HEADERS):
            cell = header_row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            self._set_run_font(run, self.config["header_font"], self.config["table_font_size_pt"])
            # 灰色背景
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "D9E2F3",
            })
            shading.append(shd)

        # 数据行
        for row_idx, item in enumerate(items):
            row = table.rows[row_idx + 1]
            values = [
                str(item.id or ""),
                item.title,
                item.source,
                item.description,
                item.proof_purpose,
                item.page_range,
                item.remarks,
            ]
            for col_idx, val in enumerate(values):
                cell = row.cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                # 序号和页码居中，其余左对齐
                if col_idx in (0, 5):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                self._set_run_font(run, self.config["body_font"], self.config["table_font_size_pt"])

        # 垂直居中所有单元格
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                v_align = tc_pr.makeelement(qn("w:vAlign"), {qn("w:val"): "center"})
                tc_pr.append(v_align)

    # ------------------------------------------------------------------ #
    #  页脚：提交人 + 日期
    # ------------------------------------------------------------------ #

    def _add_footer(self, doc: Document, case_info: dict) -> None:
        """添加页脚信息：提交人与日期。"""
        doc.add_paragraph()  # 空行

        submitter = case_info.get("plaintiff", "")
        p_submit = doc.add_paragraph()
        p_submit.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_submit.add_run(f"提交人：{submitter}")
        self._set_run_font(run, self.config["body_font"], self.config["body_size_pt"])

        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_str = datetime.now().strftime("%Y年%m月%d日")
        run = p_date.add_run(f"日　期：{date_str}")
        self._set_run_font(run, self.config["body_font"], self.config["body_size_pt"])

    # ------------------------------------------------------------------ #
    #  页码
    # ------------------------------------------------------------------ #

    @staticmethod
    def _add_page_number(doc: Document) -> None:
        """在页脚添加居中页码。"""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # PAGE 域代码
        run = p.add_run()
        fld_char_begin = run.element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run.element.append(fld_char_begin)

        run2 = p.add_run()
        instr = run2.element.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
        instr.text = " PAGE "
        run2.element.append(instr)

        run3 = p.add_run()
        fld_char_end = run3.element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run3.element.append(fld_char_end)

    # ------------------------------------------------------------------ #
    #  辅助方法
    # ------------------------------------------------------------------ #

    @staticmethod
    def _set_run_font(run, font_name: str, size_pt: float) -> None:
        """设置 run 的中西文字体和字号。"""
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
