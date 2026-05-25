"""Generate printable Word evidence documents from image batches."""

from __future__ import annotations

from datetime import datetime
from math import ceil
from pathlib import Path
import re
from shutil import rmtree
from typing import Any
from uuid import uuid4

from PIL import Image, ImageOps
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Cm, Pt

SUPPORTED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".webp"}
MAX_IMAGES = 50


def generate_image_evidence_docx(
    image_paths: list[str | Path],
    *,
    title: str = "图片证据材料",
    show_filename: bool = True,
    show_index: bool = True,
    output_root: str | Path | None = None,
    max_images: int = MAX_IMAGES,
) -> dict[str, Any]:
    """Create a 2x2-per-page A4 Word document from image paths."""

    output_dir = Path(output_root or Path.cwd() / "output" / "tools" / "image_evidence_docx")
    output_dir.mkdir(parents=True, exist_ok=True)
    work_dir = output_dir / "normalized" / datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    work_dir.mkdir(parents=True, exist_ok=True)

    valid_images, skipped = _validate_images(
        image_paths,
        max_images=max_images,
        work_dir=work_dir,
    )
    if not valid_images:
        return {
            "ok": False,
            "tool_id": "image_evidence_docx",
            "error": "没有可用图片，请选择 jpg、jpeg、png、bmp 或 webp 格式图片。",
            "file_path": "",
            "image_count": 0,
            "page_count": 0,
            "skipped": skipped,
        }

    try:
        doc = Document()
        _configure_section(doc.sections[0])

        _add_title(doc, title)

        for page_index, start in enumerate(range(0, len(valid_images), 4)):
            if page_index > 0:
                doc.add_page_break()

            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            for row in table.rows:
                row.height = Cm(11.1)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                for cell in row.cells:
                    cell.width = Cm(9)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            for offset, image_info in enumerate(valid_images[start : start + 4]):
                row_index, col_index = divmod(offset, 2)
                cell = table.cell(row_index, col_index)
                _fill_image_cell(
                    cell,
                    image_info=image_info,
                    display_index=start + offset + 1,
                    show_filename=show_filename,
                    show_index=show_index,
                )

        safe_title = _safe_filename(title or "图片证据材料")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        file_path = output_dir / f"{safe_title}_{timestamp}_{uuid4().hex[:6]}.docx"
        doc.save(file_path)
    finally:
        rmtree(work_dir, ignore_errors=True)

    return {
        "ok": True,
        "tool_id": "image_evidence_docx",
        "file_path": str(file_path),
        "image_count": len(valid_images),
        "page_count": ceil(len(valid_images) / 4),
        "skipped": skipped,
    }


def _validate_images(
    image_paths: list[str | Path],
    *,
    max_images: int,
    work_dir: Path,
) -> tuple[list[dict[str, Any]], list[dict[str, str]]]:
    valid: list[dict[str, Any]] = []
    skipped: list[dict[str, str]] = []

    for index, raw_path in enumerate(image_paths):
        path = Path(raw_path)
        if index >= max_images:
            skipped.append({"file": path.name, "reason": f"超过单次最多 {max_images} 张限制"})
            continue
        if not path.is_file():
            skipped.append({"file": str(path), "reason": "文件不存在"})
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            skipped.append({"file": path.name, "reason": "不支持的图片格式"})
            continue
        try:
            normalized_path, width, height = _normalize_image_for_word(
                path,
                work_dir=work_dir,
                display_index=len(valid) + 1,
            )
        except Exception as exc:
            skipped.append({"file": path.name, "reason": f"图片无法打开：{exc}"})
            continue
        if width <= 0 or height <= 0:
            skipped.append({"file": path.name, "reason": "图片尺寸异常"})
            continue
        valid.append({
            "path": normalized_path,
            "source_name": path.name,
            "width": width,
            "height": height,
        })

    return valid, skipped


def _normalize_image_for_word(
    path: Path,
    *,
    work_dir: Path,
    display_index: int,
) -> tuple[Path, int, int]:
    """Convert arbitrary supported images into Word-friendly PNG/JPEG files."""

    with Image.open(path) as image:
        image = ImageOps.exif_transpose(image)
        image.load()
        if image.width <= 0 or image.height <= 0:
            raise ValueError("图片尺寸异常")

        if image.mode in {"RGBA", "LA"} or (
            image.mode == "P" and "transparency" in image.info
        ):
            normalized = image.convert("RGBA")
            background = Image.new("RGBA", normalized.size, (255, 255, 255, 255))
            background.alpha_composite(normalized)
            converted = background.convert("RGB")
        else:
            converted = image.convert("RGB")

        converted.thumbnail((2400, 2400), Image.Resampling.LANCZOS)
        target = work_dir / f"{display_index:03d}_{_safe_filename(path.stem)}.jpg"
        converted.save(target, format="JPEG", quality=92, optimize=True, progressive=False)
        return target, converted.width, converted.height


def _configure_section(section: Any) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)


def _add_title(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(title or "图片证据材料")
    run.bold = True
    run.font.size = Pt(16)


def _fill_image_cell(
    cell: Any,
    *,
    image_info: dict[str, Any],
    display_index: int,
    show_filename: bool,
    show_index: bool,
) -> None:
    path: Path = image_info["path"]
    source_name = str(image_info.get("source_name") or path.name)
    width_px = float(image_info["width"])
    height_px = float(image_info["height"])
    width_cm, height_cm = _fit_size(width_px, height_px, max_width_cm=8.3, max_height_cm=9.1)

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))

    caption = _caption(
        _short_filename(source_name),
        display_index,
        show_filename=show_filename,
        show_index=show_index,
    )
    if caption:
        caption_paragraph = cell.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.paragraph_format.space_before = Pt(0)
        caption_paragraph.paragraph_format.space_after = Pt(0)
        caption_paragraph.paragraph_format.line_spacing = 1
        caption_run = caption_paragraph.add_run(caption)
        caption_run.font.size = Pt(8)


def _fit_size(
    width_px: float,
    height_px: float,
    *,
    max_width_cm: float,
    max_height_cm: float,
) -> tuple[float, float]:
    aspect = width_px / height_px
    width_cm = max_width_cm
    height_cm = width_cm / aspect
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * aspect
    return width_cm, height_cm


def _caption(
    filename: str,
    display_index: int,
    *,
    show_filename: bool,
    show_index: bool,
) -> str:
    parts: list[str] = []
    if show_index:
        parts.append(f"图{display_index}")
    if show_filename:
        parts.append(filename)
    if len(parts) == 2:
        return f"{parts[0]}：{parts[1]}"
    return parts[0] if parts else ""


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]+', "_", value.strip())
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned[:48] or "图片证据材料"


def _short_filename(value: str, max_length: int = 42) -> str:
    if len(value) <= max_length:
        return value
    path = Path(value)
    suffix = path.suffix
    stem_limit = max(12, max_length - len(suffix) - 3)
    return f"{path.stem[:stem_limit]}...{suffix}"
